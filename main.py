import subprocess
import argparse
import os
import time
from concurrent.futures import ThreadPoolExecutor
import logging
import shutil
import tempfile
import google.generativeai as genai
import traceback
import docx
from docx import Document
import re  # Импортируем модуль regular expressions
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

# Настройка логирования
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Загрузка переменных окружения из файла .env
load_dotenv()

# Получение настроек API из переменных окружения
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL = os.getenv("GEMINI_MODEL")

# Путь к папке с видео и папке для сохранения транскрипций
VIDEO_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), "video")
TRANSCRIBED_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), "transcribed_text")

# Создаем папку для транскрибированных текстов, если она не существует
os.makedirs(TRANSCRIBED_FOLDER, exist_ok=True)

# Флаг для отключения транскрипции, если API недоступен
SKIP_TRANSCRIPTION = False  # Установите True, чтобы пропустить транскрипцию и создать пустой документ

try:
    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel(GEMINI_MODEL)  # Используем модель из файла .env
    logging.info(f"API Gemini успешно инициализирован, используется модель {GEMINI_MODEL}")
except Exception as e:
    logging.error(f"Ошибка при инициализации API Gemini: {e}")
    SKIP_TRANSCRIPTION = True  # Автоматически отключаем транскрипцию при ошибке инициализации

def process_with_neural_network(audio_file):
    """
    Обработка аудиофайла нейросетью для получения транскрипции
    """
    try:
        start_time = time.time()
        logging.info(f"Начало транскрипции аудио: {audio_file}")

        with open(audio_file, 'rb') as f:
            audio_bytes = f.read()

        prompt = (
            "Decode the following audio in the language in which it was spoken. Remove filler words such as “uh,” “uh,” “um,” “ah,” etc.  Present the transcription as clear, coherent and grammatically correct text in the language spoken, without unnecessary pauses or hesitation."
        )

        response = model.generate_content([{"mime_type": f"audio/{audio_file.split('.')[-1]}", "data": audio_bytes}, prompt])

        end_time = time.time()
        logging.info(f"Транскрипция завершена за {end_time - start_time:.2f} секунд для {audio_file}")
        return response.text
    except Exception as e:
        logging.error(f"Ошибка при транскрипции {audio_file}: {e}")
        traceback.print_exc()
        return None




def save_transcription_to_docx(transcription, output_file):
    """
    Сохранить транскрипцию в файл .docx, корректно обрабатывая Markdown,
    включая форматирование внутри элементов списков и выравнивание по ширине.
    """
    doc = Document()

    lines = transcription.splitlines()

    for line in lines:
        line = line.strip()

        if line.startswith("#"):  # Заголовки
            level = line.count("#")
            if level > 9:
                level = 9  # Ограничиваем уровень заголовка
            heading_text = line.lstrip('#').strip()
            doc.add_heading(heading_text, level=level)

        elif line.startswith("- ") or line.startswith("* "):  # Маркированный список
            list_text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            apply_formatting_to_paragraph(list_text, p)

        elif re.match(r'^\d+\.\s', line):  # Нумерованный список
            list_text = re.sub(r'^\d+\.\s', '', line).strip()
            p = doc.add_paragraph(style='List Number')
            apply_formatting_to_paragraph(list_text, p)

        elif line.startswith("> "):  # Цитата
            quote_text = line[2:].strip()
            p = doc.add_paragraph(style='Quote')
            apply_formatting_to_paragraph(quote_text, p)

        elif line == "":  # Пустая строка
            doc.add_paragraph()

        else:  # Обычный текст
            p = doc.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине
            apply_formatting_to_paragraph(line, p)

    # Создаем папку для файла, если она не существует
    os.makedirs(os.path.dirname(output_file), exist_ok=True)
    doc.save(output_file)
    logging.info(f"Транскрипция сохранена в {output_file}")

def apply_formatting_to_paragraph(text, paragraph):
    """
    Применяет форматирование (жирный, курсив) к тексту внутри абзаца.
    """
    text = re.sub(r'\*\*(.*?)\*\*', r'**\1**', text)  # Жирный (**)
    text = re.sub(r'__(.*?)__', r'**\1**', text)
    text = re.sub(r'\*(.*?)\*', r'*\1*', text)  # Курсив (*)
    text = re.sub(r'_(.*?)_', r'*\1*', text)

    segments = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)

    for segment in segments:
        if segment.startswith('**') and segment.endswith('**'):
            run = paragraph.add_run(segment[2:-2])
            run.bold = True
        elif segment.startswith('*') and segment.endswith('*'):
            run = paragraph.add_run(segment[1:-1])
            run.italic = True
        else:
            paragraph.add_run(segment)

def merge_docx_files(docx_files, output_path):
    """
    Объединить несколько .docx файлов в один, сохраняя форматирование.
    """
    merged_doc = Document()

    for docx_file in docx_files:
        try:
            doc = Document(docx_file)
            for element in doc.element.body:
                merged_doc.element.body.append(element)
        except Exception as e:
            logging.error(f"Ошибка при объединении файла {docx_file}: {e}")

    # Создаем папку для файла, если она не существует
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    merged_doc.save(output_path)
    logging.info(f"Файлы успешно объединены в {output_path}")

    # Удалить все промежуточные файлы
    for docx_file in docx_files:
        os.remove(docx_file)

    logging.info("Промежуточные файлы .docx удалены")

def extract_audio_sequential(input_video, segment_duration=600, output_dir=None, audio_format='mp3', audio_bitrate='192k', max_transcription_workers=4):
    """
    Последовательное извлечение аудио из видео и его параллельная обработка нейросетью
    """
    try:
        base_name = os.path.splitext(os.path.basename(input_video))[0]
        file_extension = f".{audio_format}"
        
        # Проверяем, есть ли уже транскрибированный файл для этого видео
        transcribed_file = os.path.join(TRANSCRIBED_FOLDER, f"{base_name}.docx")
        if os.path.exists(transcribed_file):
            logging.info(f"Файл {transcribed_file} уже существует, пропускаем транскрипцию")
            return True

        # Создаем временную папку для обработки
        temp_dir = tempfile.mkdtemp()

        start_time = time.time()
        logging.info(f"Начало извлечения аудио: {input_video}")
        
        # Используем временную папку для аудио сегментов
        if output_dir is None:
            output_dir = temp_dir
        else:
            os.makedirs(output_dir, exist_ok=True)
            
        temp_audio = os.path.join(temp_dir, f"temp_audio{file_extension}")
        extract_full_audio_cmd = [
            'ffmpeg',
            '-i', input_video,
            '-vn',
            '-acodec', audio_format,
            '-ab', audio_bitrate,
            '-y',
            temp_audio
        ]

        subprocess.run(extract_full_audio_cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

        segment_count = 0
        audio_segments = []  # Список для хранения путей к аудиофайлам

        # Извлечение всех сегментов
        while True:
            segment_count += 1
            output_file = os.path.join(output_dir, f'{base_name}_audio_{segment_count:03d}{file_extension}')

            probe_cmd = [
                'ffprobe',
                '-v', 'error',
                '-show_entries', 'format=duration',
                '-of', 'default=noprint_wrappers=1:nokey=1',
                temp_audio
            ]

            try:
                duration_output = subprocess.check_output(probe_cmd).decode('utf-8').strip()
                if not duration_output or float(duration_output) < 1:
                    logging.info(f"Достигнут конец аудио после создания {segment_count-1} сегментов")
                    break
                remaining_duration = float(duration_output)
            except (subprocess.SubprocessError, ValueError) as e:
                logging.error(f"Ошибка при проверке продолжительности: {e}")
                break

            current_segment_duration = min(segment_duration, remaining_duration)

            extract_cmd = [
                'ffmpeg',
                '-i', temp_audio,
                '-t', str(current_segment_duration),
                '-acodec', 'copy',
                '-y',
                output_file
            ]

            logging.info(f"Создание аудио сегмента {segment_count} ({current_segment_duration:.2f} секунд)")
            subprocess.run(extract_cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

            audio_segments.append(output_file)

            if remaining_duration <= segment_duration:
                logging.info(f"Достигнут конец аудио после создания {segment_count} сегментов")
                break

            temp_output = os.path.join(temp_dir, f"temp_output{file_extension}")
            trim_cmd = [
                'ffmpeg',
                '-i', temp_audio,
                '-ss', str(current_segment_duration),
                '-acodec', 'copy',
                '-y',
                temp_output
            ]

            subprocess.run(trim_cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            os.remove(temp_audio)
            os.rename(temp_output, temp_audio)

        # Параллельная обработка аудиосегментов
        docx_files = []  # Список для хранения путей к файлам .docx с транскрипциями
        
        if SKIP_TRANSCRIPTION:
            # Если транскрипция отключена, создаем пустой документ
            logging.info("Транскрипция отключена. Создаем пустой документ.")
            dummy_file = os.path.join(output_dir, f"{base_name}_dummy_transcript.docx")
            doc = Document()
            doc.add_paragraph(f"Аудио извлечено из {input_video}. Транскрипция отключена из-за недоступности API.")
            doc.save(dummy_file)
            docx_files.append(dummy_file)
        else:
            # Обычная обработка с транскрипцией
            with ThreadPoolExecutor(max_workers=max_transcription_workers) as executor:
                future_to_file = {executor.submit(process_with_neural_network, audio_file): audio_file for audio_file in audio_segments}
                for future in future_to_file:
                    audio_file = future_to_file[future]
                    try:
                        transcription = future.result()
                        if transcription:
                            # Сохраняем транскрипцию напрямую, без структуризации
                            transcription_file = f"{audio_file.rsplit('.', 1)[0]}_transcript.docx"
                            save_transcription_to_docx(transcription, transcription_file)
                            docx_files.append(transcription_file)
                    except Exception as e:
                        logging.error(f"Ошибка при обработке результата транскрипции для {audio_file}: {e}")
                        
            # Если после всех попыток транскрипции список файлов пуст, создаем пустой документ
            if not docx_files:
                logging.warning("Не удалось получить ни одной транскрипции. Создаем пустой документ.")
                dummy_file = os.path.join(output_dir, f"{base_name}_empty_transcript.docx")
                doc = Document()
                doc.add_paragraph(f"Не удалось получить транскрипцию для аудио из {input_video}. Проверьте логи для подробностей.")
                doc.save(dummy_file)
                docx_files.append(dummy_file)

        # Объединение всех .docx файлов в один и сохранение в папке transcribed_text
        final_transcript_file = os.path.join(TRANSCRIBED_FOLDER, f"{base_name}.docx")
        merge_docx_files(docx_files, final_transcript_file)
        logging.info(f"Файл {final_transcript_file} успешно создан.")
        
        # Убрано ожидание и дополнительная обработка структуры

        # Удаление аудиофайлов после обработки
        for audio_segment in audio_segments:
            os.remove(audio_segment)

        shutil.rmtree(temp_dir)
        end_time = time.time()
        logging.info(f"Завершено: {input_video} (заняло {end_time - start_time:.2f} секунд)")
        return True
    except Exception as e:
        logging.error(f"Ошибка при обработке {input_video}: {e}")
        try:
            shutil.rmtree(temp_dir)
        except:
            pass
        return False


def process_videos_in_folder(segment_duration=600, output_dir=None, max_workers=None, audio_format='mp3', audio_bitrate='192k', max_transcription_workers=4):
    """Обработать все видеофайлы в папке video, извлекая аудио и транскрибируя его"""
    video_extensions = ('.mp4', '.mkv', '.avi', '.mov', '.flv', '.wmv')
    video_files = [os.path.join(VIDEO_FOLDER, filename)
                  for filename in os.listdir(VIDEO_FOLDER)
                  if filename.lower().endswith(video_extensions)]

    if not video_files:
        logging.warning(f"В папке {VIDEO_FOLDER} не найдено видеофайлов.")
        return

    logging.info(f"Найдено {len(video_files)} видеофайлов для извлечения аудио.")

    if output_dir is None:
        output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "audio_segments")

    os.makedirs(output_dir, exist_ok=True)

    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        for video_file in video_files:
            executor.submit(extract_audio_sequential, video_file, segment_duration, output_dir, audio_format, audio_bitrate, max_transcription_workers)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Автоматическая транскрипция видео из папки video")
    parser.add_argument('--output', '-o', help="Директория для сохранения аудио сегментов")
    parser.add_argument('--duration', '-d', type=int, default=600,
                        help="Длительность сегмента в секундах (по умолчанию 600 секунд = 10 минут)")
    parser.add_argument('--workers', '-w', type=int, default=None,
                        help="Максимальное количество параллельных процессов для видео (по умолчанию: количество CPU)")
    parser.add_argument('--transcription-workers', '-tw', type=int, default=10,
                        help="Максимальное количество параллельных потоков для транскрипции (по умолчанию: 10)")
    parser.add_argument('--format', '-f', type=str, default='mp3',
                        help="Формат аудио (mp3, aac, ogg, wav и т.д.)")
    parser.add_argument('--bitrate', '-b', type=str, default='192k',
                        help="Битрейт аудио (напр. 128k, 192k, 320k)")

    args = parser.parse_args()

    # Автоматически используем папку video для поиска видеофайлов
    process_videos_in_folder(args.duration, args.output, args.workers, args.format, args.bitrate, args.transcription_workers)